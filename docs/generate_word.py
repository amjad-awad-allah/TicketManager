import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_toc(paragraph):
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    run._r.append(instrText)
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'separate')
    run._r.append(fldChar)
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)

def add_code_block(doc, title, code_text):
    doc.add_heading(title, 3)
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.rows[0].cells[0]
    cell.paragraphs[0].text = code_text
    run = cell.paragraphs[0].runs[0]
    run.font.name = 'Consolas'
    run.font.size = Pt(8.5)
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'F2F2F2')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def create_documentation():
    doc = docx.Document()

    # Style definitions
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # 1. Deckblatt
    doc.add_paragraph('\n\n\n')
    title = doc.add_heading('PROJEKTDOKUMENTATION', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('\n')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Entwicklung eines objektorientierten IT-Ticket-Management-Systems\nmit REST-API Anbindung und lokaler Datenhaltung\n')
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(31, 73, 125)
    doc.add_paragraph('\n\n\n\n')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('Name: ').bold = True
    p.add_run('Amjad Awad-Allah\n')
    p.add_run('Fachrichtung: ').bold = True
    p.add_run('Fachinformatik – Anwendungsentwicklung\n')
    p.add_run('Ausbildungsstätte: ').bold = True
    p.add_run('Bfz-Essen\n')
    p.add_run('Projektumfang: ').bold = True
    p.add_run('64 Stunden\n')
    doc.add_page_break()

    # 2. Inhaltsverzeichnis
    doc.add_heading('Inhaltsverzeichnis', 1)
    p = doc.add_paragraph()
    add_toc(p)
    doc.add_paragraph('\n(Hinweis: In Word Rechtsklick -> "Felder aktualisieren" wählen)')
    doc.add_page_break()

    # 3. Projektplanung
    doc.add_heading('1. Projektplanung', 1)
    
    doc.add_heading('1.1 Beschreibung des Soll-Zustandes', 2)
    doc.add_paragraph(
        "Das Ziel des Projekts ist die Schaffung einer zentralen Desktop-Applikation zur effizienten Verwaltung von Support-Anfragen. "
        "Tickets sollen mit allen relevanten Metadaten (Titel, Beschreibung, Status, Priorität) erfasst werden können. "
        "Zudem soll eine Anbindung an eine externe Kunden-API erfolgen, um Stammdaten automatisiert zu importieren. "
        "Alle Daten müssen lokal und persistent gespeichert werden, um nach einem Neustart verfügbar zu sein."
    )

    doc.add_heading('1.2 Verwendete Ressourcen und Begründung', 2)
    doc.add_paragraph(
        "• Java 21: Gewählt wegen der Plattformunabhängigkeit, starken Typisierung und der modernen Features für skalierbare Software.\n"
        "• Swing Framework: Für die GUI-Entwicklung gewählt, da es leichtgewichtig ist und keine externen Laufzeitumgebungen für einfache Desktop-Tools erfordert.\n"
        "• Maven: Als Build-Management-Tool zur sauberen Verwaltung der Abhängigkeiten (Retrofit, Jackson).\n"
        "• Retrofit 2.9: Als HTTP-Client gewählt, da er eine typsichere Abstraktionsebene für REST-Calls bietet.\n"
        "• Java Serialisierung: Gewählt für eine schnelle und unkomplizierte Persistenz ohne die Komplexität einer vollwertigen Datenbank."
    )

    doc.add_heading('1.3 Mein Aufgabenbereich', 2)
    doc.add_paragraph(
        "Als alleiniger Entwickler umfasste mein Verantwortungsbereich den gesamten Software-Lifecycle:\n"
        "1. Anforderungsanalyse und Ist-Soll-Konzeption.\n"
        "2. Software-Architektur und UML-Modellierung.\n"
        "3. Implementierung der Backend-Logik und Daten-Persistenz.\n"
        "4. Design und Umsetzung der Benutzeroberfläche (Swing).\n"
        "5. Qualitätssicherung durch Unit-Tests und manuelle Abnahme."
    )

    doc.add_heading('1.4 Detaillierter Zeitplan (64h)', 2)
    time_table = doc.add_table(rows=1, cols=3)
    time_table.style = 'Table Grid'
    hdr = time_table.rows[0].cells
    hdr[0].text = 'Phase'; hdr[1].text = 'Details'; hdr[2].text = 'Dauer'
    for ph, det, dur in [('Analyse', 'Ist/Soll, Ressourcenwahl', '10h'), ('Entwurf', 'Architektur, UML, GUI-Design', '10h'), ('Implementation', 'Backend, API, Frontend', '32h'), ('Testen', 'JUnit, Bugfixing', '8h'), ('Abschluss', 'Dokumentation', '4h')]:
        row = time_table.add_row().cells
        row[0].text = ph; row[1].text = det; row[2].text = dur

    doc.add_page_break()

    # 4. Analyse und Wirtschaftlichkeit
    doc.add_heading('2. Analyse und Wirtschaftlichkeit', 1)
    
    doc.add_heading('2.1 Ist-Analyse', 2)
    doc.add_paragraph("Bisher: Unstrukturierte Erfassung von Tickets, keine zentrale Datenbank, Risiko von Datenverlust und fehlender Priorisierungsübersicht.")

    doc.add_heading('2.2 Kosten-Nutzen-Analyse', 2)
    doc.add_paragraph(
        "Die Nutzung von Java Serialisierung sparte ca. 15 Stunden Entwicklungszeit im Vergleich zu einer SQL-Datenbank (Wegfall von Setup/Mapping). "
        "Dieser Zeitgewinn wurde genutzt, um die API-Integration (Retrofit) robuster zu gestalten und ein modernes Dashboard zu entwickeln."
    )

    doc.add_page_break()

    # 5. Entwurf und Implementierung
    doc.add_heading('3. Implementierung', 1)
    
    doc.add_heading('3.1 UML-Modellierung', 2)
    doc.add_paragraph("Das System nutzt ein MVC-Muster. Die Klassenhierarchie (Benutzer -> Kunde/Admin) und das Generic Repository bilden den stabilen Kern.")

    add_code_block(doc, "3.2 Code: API-Import (Retrofit)",
        "public interface JsonPlaceholderApi {\n"
        "    @GET(\"users\")\n"
        "    Call<List<UserResponse>> getUsers();\n"
        "}")

    doc.add_heading('3.3 Projektabweichungen und Änderungen', 2)
    doc.add_paragraph(
        "Während der Implementierung gab es folgende Anpassungen:\n"
        "• Speichermedium: Ursprünglich war ein einfaches CSV-Format geplant. Dies wurde zugunsten der Binär-Serialisierung geändert, "
        "um Objektbeziehungen (Ticket zu Kunde) konsistenter laden und speichern zu können.\n"
        "• GUI: Zur Verbesserung der Übersicht wurde der Fokus auf ein Single-Window Dashboard mit Modal-Dialogen gelegt."
    )

    doc.add_page_break()

    # 6. Testen
    doc.add_heading('4. Testen', 1)
    doc.add_paragraph("Das Projekt wurde mittels White-Box (JUnit) und Black-Box (Manuelle GUI-Tests) validiert.")
    
    test_table = doc.add_table(rows=1, cols=4)
    test_table.style = 'Table Grid'
    hdr = test_table.rows[0].cells
    hdr[0].text = 'ID'; hdr[1].text = 'Szenario'; hdr[2].text = 'Ergebnis'; hdr[3].text = 'Status'
    for tid, cas, exp, stat in [('1', 'Ticket ohne Titel', 'Fehler-Dialog', 'PASS'), ('2', 'API-Import', 'Liste gefüllt', 'PASS'), ('3', 'Persistence-Test', 'Daten nach Neustart da', 'PASS')]:
        row = test_table.add_row().cells
        row[0].text = tid; row[1].text = cas; row[2].text = exp; row[3].text = stat

    doc.add_page_break()

    # 7. Fazit
    doc.add_heading('5. Fazit', 1)
    
    doc.add_heading('5.1 Erfolgreiche Umsetzung', 2)
    doc.add_paragraph("Alle Muss-Anforderungen (CRUD, Persistenz, API-Anbindung) wurden im zeitlichen Rahmen von 64 Stunden erfolgreich implementiert.")
    
    doc.add_heading('5.2 Erweiterbarkeit und Wiederverwendbarkeit', 2)
    doc.add_paragraph(
        "Durch das gewählte Repository-Muster und die Nutzung von Java Generics ist das System hochgradig erweiterbar. "
        "Es können problemlos neue Entitäten (z.B. Equipment oder Standorte) hinzugefügt werden, ohne die Kernarchitektur zu verändern."
    )
    
    doc.add_heading('5.3 Lessons Learned', 2)
    doc.add_paragraph("Die größte Hürde war die moderne Gestaltung einer Swing Oberfläche. Die Erfahrung im Umgang mit Graphics2D war wertvoll.")

    doc.add_page_break()

    # 8. Anhang und Quellen
    doc.add_heading('6. Anhang und Quellen', 1)
    doc.add_heading('6.1 UML-Diagramm', 2)
    doc.add_paragraph("[UML-Diagramm im Quellcode-Verzeichnis enthalten]")
    
    doc.add_heading('6.2 Quellenverzeichnis', 2)
    doc.add_paragraph("Oracle Java 21 Docs, Retrofit 2.9, Jackson Databind, JUnit 4 Framework.")

    # Save
    path = r"c:/Users/User/Desktop/git/TicketManager/docs/PROJEKTDOKUMENTATION.docx"
    doc.save(path)
    print(f"File created successfully at {path}")

if __name__ == "__main__":
    create_documentation()
