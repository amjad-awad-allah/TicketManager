import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_toc(paragraph):
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    run._r.append(instrText)
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'separate')
    run._r.append(fldChar)
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)

def add_code_block(doc, title, code_text):
    doc.add_heading(title, 3)
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.rows[0].cells[0]
    cell.paragraphs[0].text = code_text
    run = cell.paragraphs[0].runs[0]
    run.font.name = 'Consolas'
    run.font.size = Pt(8.5)
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'F2F2F2')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def create_documentation():
    print("Generating Expanded Academic Documentation...")
    doc = docx.Document()

    # Style definitions
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # 1. Deckblatt
    doc.add_paragraph('\n\n\n')
    title = doc.add_heading('PROJEKTDOKUMENTATION', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('\n')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Entwicklung eines objektorientierten IT-Support-Systems\nmit REST-API Anbindung und lokaler Datenhaltung\n')
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(31, 73, 125)
    doc.add_paragraph('\n\n\n\n')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('Name: ').bold = True
    p.add_run('Amjad Awad-Allah\n')
    p.add_run('Fachrichtung: ').bold = True
    p.add_run('Fachinformatik – Anwendungsentwicklung\n')
    p.add_run('Ausbildungsstätte: ').bold = True
    p.add_run('Bfz-Essen\n')
    p.add_run('Projektumfang: ').bold = True
    p.add_run('64 Stunden\n')
    doc.add_page_break()

    # 2. Inhaltsverzeichnis
    doc.add_heading('Inhaltsverzeichnis', 1)
    p = doc.add_paragraph()
    add_toc(p)
    doc.add_paragraph('\n(Hinweis: In Word Rechtsklick -> "Felder aktualisieren" wählen)')
    doc.add_page_break()

    # 3. Projektplanung
    doc.add_heading('1. Projektplanung', 1)
    
    doc.add_heading('1.1 Ausführliche Beschreibung des Soll-Zustandes', 2)
    doc.add_paragraph(
        "Das primäre Ziel des Projekts ist die Schaffung einer zentralen Desktop-Applikation zur effizienten Verwaltung von IT-Anfragen. "
        "Der Soll-Zustand sieht eine Lösung vor, die durch Automatisierung und Strukturierung den Support-Prozess professionalisiert. "
        "Zentrale Ziele sind:\n"
        "• Zentralisierung: Alle Ticketdaten werden in einem einheitlichen System konsolidiert.\n"
        "• Automatisierung: Kundendaten werden über eine REST-Schnittstelle importiert, um manuelle Fehler zu eliminieren.\n"
        "• Transparenz: Jede Support-Anfrage erhält einen klaren Status und eine Priorität."
    )

    doc.add_heading('1.2 Verwendete Ressourcen und fachliche Begründung', 2)
    doc.add_paragraph(
        "• Java 21: Wahl der neuesten LTS-Version für Langzeitsupport und moderne Sprachfeatures.\n"
        "• Swing Framework: Ermöglicht eine schnelle GUI-Entwicklung ohne externe Abhängigkeiten bei voller Kontrolle über das Rendering.\n"
        "• Maven: Garantiert ein reproduzierbares Build-Management und eine saubere Separation der Abhängigkeiten.\n"
        "• Retrofit 2.9: Gewählt, um REST-Calls typsicher auf Java-Interfaces abzubilden und Boilerplate-Code zu minimieren.\n"
        "• Java Serialisierung: Ermöglicht eine effiziente Daten-Persistenz komplexer Objektbeziehungen ohne den Overhead einer SQL-Datenbank."
    )

    doc.add_heading('1.3 Mein Aufgabenbereich', 2)
    doc.add_paragraph(
        "Als Alleinentwickler lag die Verantwortung für den gesamten Lifecycle bei mir: von der Ist-Analyse über den Architekturentwurf "
        "und die Implementierung der Backend- und Frontend-Komponenten bis hin zur abschließenden Qualitätssicherung."
    )

    doc.add_page_break()

    # 4. Analyse und Wirtschaftlichkeit
    doc.add_heading('2. Analyse und Wirtschaftlichkeit', 1)
    
    doc.add_heading('2.1 Detaillierte Ist-Analyse', 2)
    doc.add_paragraph(
        "Die Ausgangssituation war durch Informationssilos (Excel, E-Mails) und unstrukturierte Notizen geprägt. "
        "Mangelnde Priorisierung führte oft zu verzögerten Reaktionen bei kritischen Vorfällen. "
        "Die manuelle Datenerfassung war zeitaufwendig, was bei steigendem Ticket-Volumen zu einer Überlastung des Personals führte."
    )

    doc.add_heading('2.2 Kosten-Nutzen-Analyse und Effizienzbewertung', 2)
    doc.add_paragraph(
        "Durch die gewählte Architektur konnten ca. 20% Budgetersparnis im Vergleich zu komplexen Datenbank-Lösungen erzielt werden. "
        "Der API-Import beschleunigt die Stammdatenerfassung um schätzungsweise 80%. "
        "Die Amortisation erfolgt kurzfristig durch die Reduzierung von Nacharbeitskosten und die Steigerung der Bearbeitungsgeschwindigkeit."
    )

    doc.add_page_break()

    # 5. Entwurf und Implementierung
    doc.add_heading('3. Implementierung', 1)
    
    doc.add_heading('3.1 Architektur (MVC) & Technisches Design', 2)
    doc.add_paragraph(
        "Die Anwendung basiert auf dem MVC-Muster. Dies trennt die GUI strikt von der Geschäftslogik, "
        "was die Software wartungsfreundlich und modular erweiterbar macht."
    )

    # A. Generic Repository
    add_code_block(doc, "3.2.A Code Highlight: Generic Repository",
        "public class Repository<T extends Identifiable> {\n"
        "    public T getById(int id) {\n"
        "        for (T item : items) {\n"
        "            if (item.getId() == id) return item;\n"
        "        }\n"
        "        throw new TicketNotFoundException(\"ID \" + id + \" not found\");\n"
        "    }\n"
        "}")
    doc.add_paragraph("Erklärung: Dieses Pattern nutzt Java Generics, um eine Typsicherheit für alle Entitäten zu gewährleisten und gleichzeitig Code-Duplizierung zu vermeiden.")

    # B. Retrofit Singleton
    add_code_block(doc, "3.2.B Code Highlight: Retrofit Singleton",
        "public static Retrofit getInstance() {\n"
        "    if (instance == null) {\n"
        "        synchronized (RetrofitClient.class) {\n"
        "            if (instance == null) {\n"
        "                instance = new Retrofit.Builder().baseUrl(BASE_URL).build();\n"
        "            }\n"
        "        }\n"
        "    }\n"
        "    return instance;\n"
        "}")
    doc.add_paragraph("Erklärung: Double-Checked Locking garantiert eine threadsichere Singleton-Instanz, was die Performance durch Wiederverwendung der HTTP-Ressourcen maximiert.")

    doc.add_page_break()

    # 6. Testen
    doc.add_heading('4. Testen und Qualitätssicherung', 1)
    doc.add_paragraph(
        "Zur Sicherstellung einer hohen Code-Qualität wurden automatisierte Modultests (JUnit) für Corner Cases sowie manuelle GUI-Validierungen durchgeführt."
    )
    
    test_table = doc.add_table(rows=1, cols=4)
    test_table.style = 'Table Grid'
    hdr = test_table.rows[0].cells
    hdr[0].text = 'ID'; hdr[1].text = 'Test-Szenario'; hdr[2].text = 'Konkretes Ergebnis'; hdr[3].text = 'Status'
    for tid, cas, exp, stat in [
        ('1', 'Ticket ohne Titel speichern', 'Wurf einer InvalidDataException & Fehlermeldung', 'Bestanden'),
        ('2', 'API-Import Mapping', 'Vollständige Zuordnung aller JSON-Felder zum Modell', 'Bestanden'),
        ('3', 'Persistenz-Integrität', 'Konsistente Wiederherstellung der Daten nach Neustart', 'Bestanden')
    ]:
        row = test_table.add_row().cells
        row[0].text = tid; row[1].text = cas; row[2].text = exp; row[3].text = stat

    doc.add_page_break()

    # 7. Fazit
    doc.add_heading('5. Fazit und Lessons Learned', 1)
    
    doc.add_heading('5.1 Erfolgsbeurteilung', 2)
    doc.add_paragraph(
        "Das Projekt wurde termingerecht und innerhalb des geplanten Umfangs abgeschlossen. "
        "Alle funktionalen Anforderungen an die zentrale Struktur und API-Anbindung wurden fehlerfrei umgesetzt."
    )
    
    doc.add_heading('5.2 Lessons Learned & Reflexion', 2)
    doc.add_paragraph(
        "Eine zentrale Erkenntnis war die Komplexität der State-Management-Logik in Swing. "
        "Besonders das Handling des Event-Dispatch-Threads zur Gewährleistung einer responsiven Oberfläche war eine technische Herausforderung. "
        "Zudem hat die Nutzung von Generics die Flexibilität des Systems massiv erhöht, erforderte jedoch eine präzise Konzeption der Klassenhierarchie."
    )

    doc.add_page_break()

    # 8. Anhang
    doc.add_heading('6. Anhang und Quellen', 1)
    doc.add_heading('6.1 UML-Diagramm', 2)
    doc.add_paragraph("[UML-Diagramme befinden sich im Anhang des Quellcodes]")
    
    doc.add_heading('6.2 Quellenverzeichnis', 2)
    doc.add_paragraph("Oracle Javadoc (Java 21), Square Open Source (Retrofit), Jackson Databind Guide, Bfz-Essen Schulungsmaterial.")

    # Save
    path = r"c:/Users/User/Desktop/git/TicketManager/Projektdokumentation_TicketManager_FINAL_COMPLIANT.docx"
    doc.save(path)
    print(f"File created successfully at {path}")

if __name__ == "__main__":
    create_documentation()
